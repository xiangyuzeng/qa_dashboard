#!/usr/bin/env python3
"""
Word (.docx) export — a bilingual monthly report mirroring the 7-section structure of the
Excel export, authored from scratch with python-docx. Reuses the column arrays / headers /
servable-filter / risk-mix logic from export_xlsx.py (single source of truth) so the two
exports never drift. Risk-Level cells are shaded by level (5-level palette) like the workbook.

Word is narrower than Excel, so each module renders a CURATED, readable column subset (the
full 12/16/16/23-column tables live in the .xlsx system-of-record). The café table keeps
门店编号 Establishment ID. Sentiment stays summary-only (no own section), same as the xlsx.

Sections: 月度摘要 Monthly Summary · 食品安全主表 · 进口出口监管 · 州地方法规 ·
咖啡馆检查 · 数据源日志 · 字段说明.
Output is a STATIC file the dashboard links to (public/exports/monthly_report.docx).

Run: npm run prep:export   (python3 prep/export_docx.py)   ·   needs: pip install python-docx
"""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Reuse the data contract + helpers from the Excel exporter (no drift, no openpyxl styling).
from export_xlsx import (
    MODULE_LABEL,
    STATUS_LABEL,
    RISK_LEVELS,
    load,
    servable,
    fmt,
    risk_mix,
    OUT_DIR,
)

OUT = os.path.join(OUT_DIR, "monthly_report.docx")

NAVY = "1F3A5F"
HIGH_RED = RGBColor(0xC0, 0x00, 0x00)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
# Risk-level cell shading — mirror export_xlsx RISK_FILL fgColors.
RISK_SHADE = {
    "高风险": "F4CCCC", "中风险": "FCE4D6", "低风险": "E2EFDA",
    "信息参考": "E7EBF0", "关注": "CFE8EF",
}
RISK_DESC = {
    "高风险": "直接影响 high — direct impact", "中风险": "中等影响 medium",
    "低风险": "较低影响 low", "信息参考": "信息参考 informational",
    "关注": "暂无直接影响，但存在复发/恶化趋势/升级可能 watch",
}

# Curated readable column subsets per module: (record-key, "中文 English" header).
FOOD_COLS = [("no", "序号 No."), ("category", "类别 Category"), ("title", "标题 Title"),
             ("source", "来源 Source"), ("publicationDate", "发布日期 Date"), ("riskLevel", "风险 Risk")]
IMPORT_COLS = [("no", "序号 No."), ("title", "标题 Title"), ("agency", "机构 Agency"),
               ("countryRegion", "国家/地区 Country"), ("regulatoryAction", "动作 Action"), ("riskLevel", "风险 Risk")]
REG_COLS = [("no", "序号 No."), ("jurisdiction", "地区 Jurisdiction"), ("regulationBillName", "法规/法案 Bill"),
            ("status", "状态 Status"), ("effectiveDate", "生效 Effective"), ("riskLevel", "风险 Risk")]
INSP_COLS = [("no", "序号 No."), ("jurisdiction", "地区 Jurisdiction"), ("brand", "品牌 Brand"),
             ("storeName", "门店 Store"), ("establishmentId", "门店编号 Establishment ID"),
             ("inspectionResult", "结果 Result"), ("riskLevel", "风险 Risk")]


# ── low-level docx helpers ──
def _shade(cell, hex_color):
    """Set a table cell's background fill (w:shd) — no high-level python-docx API for this."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _repeat_header(row):
    """Mark a table row to repeat as a header on each printed page."""
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def _set_cell(cell, text, *, size=8, bold=False, color=None, align_center=False):
    cell.text = "" if text is None else str(text)
    for p in cell.paragraphs:
        if align_center:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if not p.runs:
            p.add_run("")
        for run in p.runs:
            run.font.size = Pt(size)
            run.font.bold = bold
            if color is not None:
                run.font.color.rgb = color


def _cell_value(rec, key):
    if key == "title":
        parts = [p for p in [rec.get("chineseTitle"), rec.get("englishTitle")] if p]
        return "  /  ".join(parts) if parts else "—"
    v = fmt(key, rec.get(key))
    return str(v) if v not in (None, "") else "—"


def add_module_table(doc, cols, records):
    if not records:
        doc.add_paragraph("（本期无数据 — none this period）")
        return
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    _repeat_header(hdr)
    for ci, (_key, label) in enumerate(cols):
        _shade(hdr.cells[ci], NAVY)
        _set_cell(hdr.cells[ci], label, size=7.5, bold=True, color=WHITE)
    for i, rec in enumerate(records):
        row = table.add_row()
        high = rec.get("riskLevel") == "高风险"
        for ci, (key, _label) in enumerate(cols):
            val = i + 1 if key == "no" else _cell_value(rec, key)
            _set_cell(row.cells[ci], val, size=7.5, bold=high,
                      color=HIGH_RED if high else None, align_center=key in ("no", "riskLevel"))
            if key == "riskLevel":
                lvl = rec.get("riskLevel")
                if lvl in RISK_SHADE:
                    _shade(row.cells[ci], RISK_SHADE[lvl])
            elif high:
                _shade(row.cells[ci], RISK_SHADE["高风险"])


def _bilingual(zh, en):
    parts = [p for p in [zh, en] if p]
    return "  /  ".join(parts)


def add_bullets(doc, items):
    for it in items:
        line = _bilingual(it.get("zh"), it.get("en"))
        if line:
            doc.add_paragraph(line, style="List Bullet")


# ── document sections ──
def build_summary(doc, meta, matrix):
    summary = meta.get("summary") or {}
    counts = meta.get("counts", {})
    period = meta.get("reportingPeriod", {})

    title = doc.add_heading(summary.get("reportNameZh", "月度摘要 Monthly Summary"), level=0)
    if summary.get("reportNameEn"):
        sub = doc.add_paragraph(summary["reportNameEn"])
        sub.runs[0].font.italic = True
        sub.runs[0].font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
    doc.add_paragraph(f"数据截止 Data as of: {meta.get('dataAsOf','')}   ·   "
                      f"周期 Period: {period.get('label','')}")

    doc.add_heading("监测范围 Scope", level=2)
    scope = _bilingual(summary.get("scopeZh"), summary.get("scopeEn"))
    if scope:
        doc.add_paragraph(scope)

    if summary.get("exclusions"):
        doc.add_heading("排除项 Exclusions", level=2)
        add_bullets(doc, summary["exclusions"])

    # KPI table
    doc.add_heading("关键指标 KPIs", level=2)
    kpi = [("食品安全主表 Food Safety", counts.get("regulatory", 0)),
           ("进口出口 Import", counts.get("importExport", 0)),
           ("州地方法规 Regulation", counts.get("regulation", 0)),
           ("咖啡馆检查 Inspections", counts.get("inspections", 0)),
           ("负面舆情 Sentiment", counts.get("sentiment", 0)),
           ("高风险 High-risk", counts.get("highRisk", 0)),
           ("关注 Watch", counts.get("watch", 0))]
    kt = doc.add_table(rows=2, cols=len(kpi))
    kt.style = "Table Grid"
    for ci, (label, val) in enumerate(kpi):
        _shade(kt.rows[0].cells[ci], NAVY)
        _set_cell(kt.rows[0].cells[ci], label, size=7.5, bold=True, color=WHITE, align_center=True)
        _set_cell(kt.rows[1].cells[ci], val, size=10, bold=True,
                  color=HIGH_RED if ci == 5 else None, align_center=True)

    # risk-mix matrix (module × 5-level + total)
    doc.add_heading("风险分布 Risk Mix（模块 × 风险等级 module × risk）", level=2)
    mt = doc.add_table(rows=1, cols=len(RISK_LEVELS) + 2)
    mt.style = "Table Grid"
    head = mt.rows[0]
    _shade(head.cells[0], NAVY)
    _set_cell(head.cells[0], "模块 Module", size=7.5, bold=True, color=WHITE)
    for ci, lvl in enumerate(RISK_LEVELS):
        _shade(head.cells[ci + 1], NAVY)
        _set_cell(head.cells[ci + 1], lvl, size=7.5, bold=True, color=WHITE, align_center=True)
    _shade(head.cells[-1], NAVY)
    _set_cell(head.cells[-1], "合计 Total", size=7.5, bold=True, color=WHITE, align_center=True)
    for m in ["food_safety", "import", "regulation", "inspection", "sentiment"]:
        row = mt.add_row()
        _set_cell(row.cells[0], MODULE_LABEL[m], size=8, bold=True)
        total = 0
        for ci, lvl in enumerate(RISK_LEVELS):
            v = matrix.get(m, {}).get(lvl, 0)
            total += v
            _set_cell(row.cells[ci + 1], v, size=8, align_center=True)
            if v and lvl in RISK_SHADE:
                _shade(row.cells[ci + 1], RISK_SHADE[lvl])
        _set_cell(row.cells[-1], total, size=8, bold=True, align_center=True)

    if summary.get("keyHighlights"):
        doc.add_heading("关键要点 Key Highlights", level=2)
        add_bullets(doc, summary["keyHighlights"])

    # high-risk items table
    if summary.get("highRiskItems"):
        doc.add_heading("高风险事项 High-Risk Items", level=2)
        ht = doc.add_table(rows=1, cols=4)
        ht.style = "Table Grid"
        for ci, label in enumerate(["模块 Module", "标题 Title", "风险 Risk", "链接 Ref"]):
            _shade(ht.rows[0].cells[ci], NAVY)
            _set_cell(ht.rows[0].cells[ci], label, size=7.5, bold=True, color=WHITE)
        for it in summary["highRiskItems"]:
            row = ht.add_row()
            _set_cell(row.cells[0], MODULE_LABEL.get(it.get("module"), it.get("module", "")), size=8)
            _set_cell(row.cells[1], it.get("titleZh") or it.get("titleEn") or it.get("recordId", ""), size=8)
            lvl = it.get("riskLevel", "")
            _set_cell(row.cells[2], lvl, size=8, bold=(lvl == "高风险"),
                      color=HIGH_RED if lvl == "高风险" else None, align_center=True)
            if lvl in RISK_SHADE:
                _shade(row.cells[2], RISK_SHADE[lvl])
            _set_cell(row.cells[3], it.get("href", ""), size=8)

    if summary.get("keyActions"):
        doc.add_heading("关键行动 Key Actions", level=2)
        add_bullets(doc, summary["keyActions"])


def build_sources(doc, meta):
    doc.add_heading("数据源日志 Data Sources & Pull Log", level=1)
    prov = meta.get("provenance", [])
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    for ci, label in enumerate(["数据源 Source", "模块 Module", "状态 Status", "记录数 Records", "原文链接 URL", "备注 Notes"]):
        _shade(table.rows[0].cells[ci], NAVY)
        _set_cell(table.rows[0].cells[ci], label, size=7.5, bold=True, color=WHITE)
    for p in prov:
        row = table.add_row()
        _set_cell(row.cells[0], p.get("name", ""), size=7.5)
        _set_cell(row.cells[1], MODULE_LABEL.get(p.get("module"), p.get("module", "")), size=7.5)
        _set_cell(row.cells[2], STATUS_LABEL.get(p.get("status"), p.get("status", "")), size=7.5)
        _set_cell(row.cells[3], p.get("recordCount", 0), size=7.5, align_center=True)
        _set_cell(row.cells[4], p.get("endpointOrUrl") or "", size=7)
        _set_cell(row.cells[5], p.get("stalenessNote") or "", size=7)


def build_field_guide(doc):
    doc.add_heading("字段说明 Field Guide", level=1)
    doc.add_paragraph("风险等级 Risk Level — 5 级，单元格按等级着色 / 5-level, cells shaded by level:")
    rt = doc.add_table(rows=1, cols=2)
    rt.style = "Table Grid"
    for ci, label in enumerate(["等级 Level", "说明 Description"]):
        _shade(rt.rows[0].cells[ci], NAVY)
        _set_cell(rt.rows[0].cells[ci], label, size=7.5, bold=True, color=WHITE)
    for lvl in RISK_LEVELS:
        row = rt.add_row()
        _set_cell(row.cells[0], lvl, size=8, bold=(lvl == "高风险"), color=HIGH_RED if lvl == "高风险" else None)
        _shade(row.cells[0], RISK_SHADE[lvl])
        _set_cell(row.cells[1], RISK_DESC[lvl], size=8)
    doc.add_paragraph("数据源状态 Status：" + " · ".join(STATUS_LABEL.values()))
    doc.add_paragraph("门店编号 Establishment ID：NYC camis / Boston licenseno / 设施编号；重复违规检测的稳定关联键。")
    doc.add_paragraph("负面舆情 Sentiment：仅汇总于月度摘要（KPI + 风险分布矩阵），明细见仪表盘 /sentiment，绝不转载文章正文。")


def verify(path):
    doc = Document(path)
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"]
    assert any("数据源日志" in h for h in headings), f"missing Sources section; headings={headings[:8]}"
    assert any("字段说明" in h for h in headings), "missing Field Guide section"
    # at least one table has a 门店编号 (Establishment ID) header cell
    has_estid = any(
        any("门店编号" in c.text for c in t.rows[0].cells)
        for t in doc.tables if t.rows
    )
    assert has_estid, "no 门店编号 column in any table"
    # at least one data-bearing table (header + ≥1 row)
    assert any(len(t.rows) >= 2 for t in doc.tables), "no populated tables"
    print(f"verify OK: {len(doc.tables)} tables, {len(headings)} section headings")


def main():
    reg = [r for r in load("regulatory.json") if servable(r)]
    insp = [r for r in load("inspections.json") if servable(r)]
    imp = [r for r in load("import_export.json") if servable(r)]
    regs = [r for r in load("regulations.json") if servable(r)]
    sent = [r for r in load("sentiment.json") if servable(r) and not r.get("excluded")]
    meta = load("meta.json")

    doc = Document()
    # base font keeps Latin tidy; CJK falls back to the viewer's font (Unicode is written fine).
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(9)

    build_summary(doc, meta, risk_mix(reg, imp, regs, insp, sent))

    doc.add_page_break()
    doc.add_heading("食品安全主表 Food Safety", level=1)
    add_module_table(doc, FOOD_COLS, reg)
    doc.add_heading("进口出口监管 Import / Export & Border Control", level=1)
    add_module_table(doc, IMPORT_COLS, imp)
    doc.add_heading("州地方法规 State & Local Regulation", level=1)
    add_module_table(doc, REG_COLS, regs)
    doc.add_heading("咖啡馆检查 Café Inspections", level=1)
    add_module_table(doc, INSP_COLS, insp)

    doc.add_page_break()
    build_sources(doc, meta)
    build_field_guide(doc)

    os.makedirs(OUT_DIR, exist_ok=True)
    doc.save(OUT)
    print(f"export(docx): summary · {len(reg)} food-safety · {len(imp)} import · {len(regs)} regulation · "
          f"{len(insp)} inspection · {len(meta.get('provenance', []))} sources → public/exports/monthly_report.docx")
    verify(OUT)


if __name__ == "__main__":
    main()
